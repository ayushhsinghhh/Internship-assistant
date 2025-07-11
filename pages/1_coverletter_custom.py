import streamlit as st
from docx import Document
import os

st.set_page_config(page_title="Cover Letter Customizer", layout="centered")
st.title("📄 Cover Letter Customizer")

st.markdown("Upload your custom cover letter or choose from predefined templates.")

# --- Option 1: Upload your own template ---
uploaded_file = st.file_uploader("📤 Upload a custom .docx cover letter", type=["docx"])
if uploaded_file:
    doc = Document(uploaded_file)
    st.subheader("🔍 Preview of uploaded letter:")
    for para in doc.paragraphs:
        st.write(para.text)

# --- Option 2: Choose from predefined templates with thumbnails ---
st.subheader("📚 Choose a sample template")

templates = {
    "General AI Internship": {
        "file": "templates/CoverLetter.docx",
        "image": "templates/thumbnails/CoverLetter.png"
    },
    "Research Role (CS)": {
        "file": "templates/email.docx",
        "image": "templates/thumbnails/email.png"
    }
}

cols = st.columns(len(templates))

for i, (name, data) in enumerate(templates.items()):
    with cols[i]:
        st.image(data["image"], caption=name, use_container_width=True)
        if st.button(f"Use Template: {name}"):
            path = data["file"]
            if os.path.exists(path):
                doc = Document(path)
                st.subheader(f"📄 Preview: {name}")
                for para in doc.paragraphs:
                    st.write(para.text)
                st.session_state["selected_template"] = path
            else:
                st.error(f"Template '{name}' not found.")
